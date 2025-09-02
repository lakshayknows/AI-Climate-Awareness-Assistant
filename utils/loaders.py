# utils/loaders.py

import os
from typing import Union
from PyPDF2 import PdfReader
import docx

def load_from_upload(uploaded_file) -> str:
    """
    Load text from an uploaded file (.txt, .pdf, .docx).
    """
    filename = uploaded_file.name.lower()

    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    elif filename.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()

    elif filename.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        raise ValueError("Unsupported file type. Please upload a .txt, .pdf, or .docx file.")


def load_from_path(file_path: str) -> str:
    """
    Load text directly from a file path.
    """
    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()

    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        raise ValueError("Unsupported file type. Only .txt, .pdf, .docx are supported.")
